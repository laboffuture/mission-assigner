"""Read SME documents from input/ into a flat list of sections.

Each section is a dict: { source_file, heading, heading_path, body }.
heading_path is a breadcrumb built from the document's heading hierarchy
(e.g. "Loops > Counting loops"). It is derived from CONTENT, never from
position, so inserting text early in a document does not renumber later
sections.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED = {".docx", ".pdf", ".md", ".txt"}
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def _push_heading(stack: list[tuple[int, str]], level: int, title: str) -> str:
    """Update the heading stack and return the breadcrumb heading_path."""
    while stack and stack[-1][0] >= level:
        stack.pop()
    stack.append((level, title))
    return " > ".join(t for _, t in stack)


def _split_markdown(text: str, source_file: str) -> list[dict]:
    sections: list[dict] = []
    stack: list[tuple[int, str]] = []
    cur_heading = "(preamble)"
    cur_path = "(preamble)"
    buf: list[str] = []

    def flush():
        body = "\n".join(buf).strip()
        if body:
            sections.append(
                {
                    "source_file": source_file,
                    "heading": cur_heading,
                    "heading_path": cur_path,
                    "body": body,
                }
            )

    for line in text.splitlines():
        m = _HEADING_RE.match(line.strip())
        if m:
            flush()
            buf = []
            level = len(m.group(1))
            cur_heading = m.group(2).strip()
            cur_path = _push_heading(stack, level, cur_heading)
        else:
            buf.append(line)
    flush()
    return sections


def _split_plaintext_blocks(text: str, source_file: str) -> list[dict]:
    """Fallback for documents with no heading structure: split on blank-line
    paragraph groups and use each block's first line as its heading (content-
    derived, so it stays stable when other blocks change)."""
    sections: list[dict] = []
    blocks = re.split(r"\n\s*\n", text)
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        first_line = block.splitlines()[0].strip()
        heading = (first_line[:80]).strip() or "(untitled)"
        sections.append(
            {
                "source_file": source_file,
                "heading": heading,
                "heading_path": heading,
                "body": block,
            }
        )
    return sections


def _read_docx(path: Path) -> list[dict]:
    doc = Document(str(path))
    heading_styles = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}
    has_headings = any(
        (p.style and p.style.name in heading_styles) for p in doc.paragraphs
    )
    if not has_headings:
        print(
            f"  WARNING: {path.name} has no Word heading styles (Heading 1/2/3). "
            f"Falling back to blank-line paragraph groups."
        )
        full = "\n\n".join(p.text for p in doc.paragraphs)
        return _split_plaintext_blocks(full, path.name)

    sections: list[dict] = []
    stack: list[tuple[int, str]] = []
    cur_heading = "(preamble)"
    cur_path = "(preamble)"
    buf: list[str] = []

    def flush():
        body = "\n".join(buf).strip()
        if body:
            sections.append(
                {
                    "source_file": path.name,
                    "heading": cur_heading,
                    "heading_path": cur_path,
                    "body": body,
                }
            )

    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style in heading_styles and p.text.strip():
            flush()
            buf = []
            cur_heading = p.text.strip()
            cur_path = _push_heading(stack, heading_styles[style], cur_heading)
        elif p.text.strip():
            buf.append(p.text)
    flush()
    return sections


def _read_pdf(path: Path) -> list[dict]:
    reader = PdfReader(str(path))
    text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
    print(
        f"  NOTE: {path.name} is a PDF; heading styles are not available in PDF text. "
        f"Splitting on blank-line paragraph groups."
    )
    return _split_plaintext_blocks(text, path.name)


def read_document(path: Path) -> list[dict]:
    ext = path.suffix.lower()
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext in (".md", ".txt"):
        text = path.read_text(encoding="utf-8")
        # .md gets heading parsing; .txt tries headings then falls back.
        sections = _split_markdown(text, path.name)
        # If markdown parsing found only the preamble (no real headings), fall back.
        if len(sections) <= 1 and not any(_HEADING_RE.match(l.strip()) for l in text.splitlines()):
            print(
                f"  WARNING: {path.name} has no markdown headings. "
                f"Falling back to blank-line paragraph groups."
            )
            return _split_plaintext_blocks(text, path.name)
        return sections
    raise ValueError(f"Unsupported file type: {path.name}")


def read_input_dir(input_dir: Path) -> list[dict]:
    """Read every supported document in input_dir. Returns a flat section list."""
    sections: list[dict] = []
    files = sorted(p for p in input_dir.iterdir() if p.suffix.lower() in SUPPORTED)
    if not files:
        print(f"  No supported documents found in {input_dir}")
        return sections
    for path in files:
        print(f"  Reading {path.name} ...")
        sections.extend(read_document(path))
    return sections
